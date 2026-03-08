#!/usr/bin/env python3
"""
Generate a meeting notes Word document from structured JSON data.

Usage:
    python create_docx.py --data '<json>' --output 'meeting-notes-2024-01-15.docx'
"""

import argparse
import json
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx is not installed. Run: pip install python-docx")
    sys.exit(1)


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_bullet(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = para.add_run(bold_prefix)
        run.bold = True
        para.add_run(text)
    else:
        para.add_run(text)
    return para


def build_document(data: dict, output_path: str):
    doc = Document()

    # Title
    meeting_date = data.get("meeting_date", str(date.today()))
    title = doc.add_heading(f"Meeting Notes — {meeting_date}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 1. Attendees
    add_heading(doc, "Attendees", level=1)
    attendees = data.get("attendees", [])
    if attendees:
        for person in attendees:
            name = person.get("name", "Unknown")
            present = person.get("present", True)
            label = name if present else f"{name} (not present)"
            add_bullet(doc, label)
    else:
        doc.add_paragraph("No attendees listed.")

    doc.add_paragraph()

    # 2. Meeting Purpose
    add_heading(doc, "Meeting Purpose", level=1)
    purpose = data.get("meeting_purpose", "")
    doc.add_paragraph(purpose if purpose else "Not specified.")

    doc.add_paragraph()

    # 3. Key Decisions
    add_heading(doc, "Key Decisions", level=1)
    decisions = data.get("key_decisions", [])
    if decisions:
        for decision in decisions:
            add_bullet(doc, decision)
    else:
        doc.add_paragraph("No explicit decisions recorded.")

    doc.add_paragraph()

    # 4. Discussion Summary
    add_heading(doc, "Discussion Summary", level=1)
    summary = data.get("discussion_summary", "")
    doc.add_paragraph(summary if summary else "No summary available.")

    doc.add_paragraph()

    # 5. Outcomes
    add_heading(doc, "Outcomes", level=1)
    outcomes = data.get("outcomes", [])
    if outcomes:
        for outcome in outcomes:
            add_bullet(doc, outcome)
    else:
        doc.add_paragraph("No outcomes recorded.")

    doc.add_paragraph()

    # 6. Action Items
    add_heading(doc, "Action Items", level=1)
    action_items = data.get("action_items", [])
    if action_items:
        for item in action_items:
            owner = item.get("owner", "Unassigned")
            action = item.get("action", "")
            deadline = item.get("deadline", "")

            para = doc.add_paragraph(style="List Bullet")
            para.add_run(f"[{owner}] ").bold = True
            para.add_run(action)
            if deadline:
                run = para.add_run(f" — by: {deadline}")
                run.italic = True
    else:
        doc.add_paragraph("No action items recorded.")

    # Save
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"Saved: {output.resolve()}")


def main():
    parser = argparse.ArgumentParser(description="Generate meeting notes DOCX")
    parser.add_argument("--data", required=True, help="JSON string with meeting data")
    parser.add_argument(
        "--output",
        default=f"meeting-notes-{date.today()}.docx",
        help="Output file path",
    )
    args = parser.parse_args()

    try:
        data = json.loads(args.data)
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON: {e}")
        sys.exit(1)

    build_document(data, args.output)


if __name__ == "__main__":
    main()
